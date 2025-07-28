
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import json
import uuid
import tempfile
import os
import re
from typing import Optional
from google import genai
from google.genai import types
from docx import Document
import io
from dotenv import load_dotenv
import firebase_admin
from firebase_admin import credentials, firestore

# Load environment variables from .env file
load_dotenv()

# Initialize Firebase Admin SDK
db = None
try:
    cred = credentials.ApplicationDefault()
    firebase_admin.initialize_app(cred)
    db = firestore.client()
except Exception as e:
    print(f"Could not initialize Firebase Admin SDK: {e}")
    print("Please ensure you have authenticated with `gcloud auth application-default login`")
    print("Running without Firebase functionality")

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Initialize Gemini AI (new SDK)
client = genai.Client()

# Load attorneys data
try:
    with open('attorneys_data.json', 'r') as f:
        USERS_DB = json.load(f)
except Exception as e:
    print(f"Error loading attorneys_data.json: {e}")
    USERS_DB = []

def get_all_specialties(db_users):
    """Extracts a unique, sorted list of all specialties from the user database."""
    all_specs = set()
    for user in db_users:
        specs = user.get("specialties", [])
        if isinstance(specs, list):
            all_specs.update(spec for spec in specs)
        elif isinstance(specs, str):
            all_specs.add(specs)
    return sorted(list(all_specs))

ALL_SPECIALTIES = get_all_specialties(USERS_DB)

# State mapping
STATE_MAP = {
    "al": "alabama", "ak": "alaska", "az": "arizona", "ar": "arkansas", "ca": "california",
    "co": "colorado", "ct": "connecticut", "de": "delaware", "fl": "florida", "ga": "georgia",
    "hi": "hawaii", "id": "idaho", "il": "illinois", "in": "indiana", "ia": "iowa",
    "ks": "kansas", "ky": "kentucky", "la": "louisiana", "me": "maine", "md": "maryland",
    "ma": "massachusetts", "mi": "michigan", "mn": "minnesota", "ms": "mississippi",
    "mo": "missouri", "mt": "montana", "ne": "nebraska", "nv": "nevada", "nh": "new hampshire",
    "nj": "new jersey", "nm": "new mexico", "ny": "new york", "nc": "north carolina",
    "nd": "north dakota", "oh": "ohio", "ok": "oklahoma", "or": "oregon", "pa": "pennsylvania",
    "ri": "rhode island", "sc": "south carolina", "sd": "south dakota", "tn": "tennessee",
    "tx": "texas", "ut": "utah", "vt": "vermont", "va": "virginia", "wa": "washington",
    "wv": "west virginia", "wi": "wisconsin", "wy": "wyoming"
}

# Specialty fallback mapping
SPECIALTY_FALLBACK_MAP = {
    "divorce": ["family law"], "child custody": ["family law"], "child support": ["family law"],
    "adoption": ["family law"], "misdemeanors": ["criminal law"], "felonies": ["criminal law"],
    "juvenile offenses": ["criminal law"], "traffic violations": ["traffic law"],
    "auto accidents": ["personal injury"], "medical malpractice": ["personal injury"],
    "business formation": ["business law", "corporate law"], "mergers & acquisitions": ["business law", "corporate law"],
    "contract disputes": ["business law", "civil litigation"]
}

def generate_goodbye_message(history: list):
    """Generate a personalized goodbye message using LLM"""
    conversation_context = ""
    if len(history) > 1:
        # Get the last few exchanges for context
        recent_messages = history[-6:]  # Last 6 messages (3 exchanges)
        conversation_context = "\n".join([
            f"{msg['role']}: {msg['parts'][0]['text']}" 
            for msg in recent_messages
        ])
    
    prompt = f"""
You are a professional legal assistant for OpenLaw. The user is saying goodbye. Generate a warm, professional, and personalized goodbye message based on the conversation context.

Guidelines:
- Be warm and appreciative of their time
- Mention that you're here if they need further assistance
- Reference any specific legal matters they discussed (if any)
- Keep it professional but friendly
- End with a positive note
- Keep it concise (2-3 sentences)

Conversation context:
{conversation_context}

Generate a natural, conversational goodbye message:
"""
    
    try:
        response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        return response.text.strip()
    except Exception as e:
        return "Thank you for using OpenLaw! If you need any further assistance with legal matters, please don't hesitate to reach out. Take care!"

def parse_query_with_gemini(history: list):
    """Parse user query with Gemini AI to extract intent and filters"""
    specialties_list = ", ".join(ALL_SPECIALTIES)
    nl_query = history[-1]['parts'][0]['text']

    prompt = f"""
You support both English and Spanish. You are extracting filters from a legal query. Use only the following options:
- specialties: Choose one or more from this dynamic list: {specialties_list}
- meetingTypes: "Virtual" or "In-person"
- hasCalendarConnected: true or false
- firm: name of the law firm
- languages: list of languages spoken
- review_keywords: a list of descriptive words from the user query (e.g., "patient", "aggressive", "responsive")
- licenseState: full U.S. state name only; We only support Texas, Florida, California, and Arizona, any other state mentioned, pass: {{"intent": "out-of-area"}}
- If query is in English, "languages": English, or if its in Spanish, "languages": Spanish
- rating: a number between 1 and 5 (e.g., 4.5 means "4.5 stars and above"). If user says things like "highly rated", "4 or above", or "only 5-star", return a number like 4.0, 4.5, or 5
- If you think user only typed a name of a person, parse it as "name": the full name, or "firstName": the first name or "lastName": the last name
- location: can be a city, ZIP code, or full state name (e.g., "Miami", "90210", "Florida"). Use only city name or ZIP when state is not mentioned.

Return a plain JSON object. Do NOT format as code, markdown, or include backticks.  

-- If the user asks for a lawyer "near me", "nearby", or similar, respond with: {{"intent": "near_me"}}
-- If the query is just a legal question or not about finding a lawyer, respond with: {{"intent": "general_question", "response": "<brief 2-5 sentence answer to the question>"}}
-- If user prompt is a greeting of some kind: {{"intent": "greeting"}}
-- If question is too random or out of legal scope, respond with: {{"intent": "out-of-scope", "witty_response": "Please provide a fun/witty reply that redirects the conversation back to legal matters/questions/finding lawyers."}}
-- If the user says that they recieved a letter from OpenLaw or something along the same lines: {{"intent": "got-letter", "letter-flow": "we follow a workflow, ask the user for a reference number. mention that the reference number is the code after the URL (https://olaw.io/REFERENCE NUMBER). after user submits reference number, ask the user "What's the ideal outcome you want to achieve for your case?, ask the user what is the level of urgency. then the preferred language: english or spanish. then ask for contact details like full name, email, phone." "}}
-- If the user wants to understand, analyze, or get help with a legal document: {{"intent": "document_help"}}
-- If the user wants to draft, create, or write a legal document: {{"intent": "draft_document"}}
-- If the user says goodbye, bye, see you, take care, or similar farewell messages: {{"intent": "goodbye"}}
-- Your name is "Ola", so if user asks your name, give a fun/witty response and them how you can help them.

Here is the full conversation history for context:
{json.dumps(history, indent=2)}

Query: "{nl_query}"
"""
    try:
        response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        cleaned_text = response.text.strip().strip("`")
        if cleaned_text.startswith("json"):
            cleaned_text = cleaned_text[4:].strip()
        return json.loads(cleaned_text)
    except (json.JSONDecodeError, Exception) as e:
        return {"error": str(e), "raw_response": response.text if 'response' in locals() else "No response"}

def find_best_matches(filters):
    """Find the best lawyer matches by applying hard filters, then scoring and ranking"""
    filtered_users = _apply_hard_filters(filters, USERS_DB)
    scored_users = [(user, _score_user(user, filters)) for user in filtered_users]
    if filters.get("location"):
        scored_users = [(user, score) for user, score in scored_users if score > 0]
    sorted_users = sorted(scored_users, key=lambda x: (x[1], x[0].get("rating", 0)), reverse=True)
    return [user for user, score in sorted_users[:3]]

def _apply_hard_filters(filters, db):
    """Apply non-negotiable filters to the user list"""
    users = db
    
    if filters.get("licenseState"):
        filter_state = filters["licenseState"].lower()
        users = [u for u in users if filter_state in [str(s).lower() for s in u.get("licenseState", [])]]

    if filters.get("specialties"):
        filter_specs = filters["specialties"]
        if not isinstance(filter_specs, list):
            filter_specs = [filter_specs]
        filter_specs = [s.lower() for s in filter_specs if s]
        if filter_specs:
            users = [u for u in users if any(spec in [s.lower() for s in u.get("specialties", [])] for spec in filter_specs)]

    if filters.get("hasCalendarConnected"):
        users = [u for u in users if u.get("hasCalendarConnected")]

    if filters.get("meetingTypes"):
        filter_meeting_types = filters["meetingTypes"]
        if not isinstance(filter_meeting_types, list):
            filter_meeting_types = [filter_meeting_types]
        filter_meeting_types = [m.lower() for m in filter_meeting_types if m]
        if filter_meeting_types:
            users = [u for u in users if any(m_type in [m.lower() for m in u.get("meetingTypes", [])] for m_type in filter_meeting_types)]

    if filters.get("firm"):
        filter_firm = filters["firm"].lower()
        users = [u for u in users if filter_firm in u.get("firm", "").lower()]

    if filters.get("languages"):
        filter_languages = filters["languages"]
        if not isinstance(filter_languages, list):
            filter_languages = [filter_languages]
        filter_languages = [language.lower() for language in filter_languages if language]
        if filter_languages:
            users = [u for u in users if any(language in [user_language.lower() for user_language in u.get("languages", [])] for language in filter_languages)]

    return users

def _score_user(user, filters):
    """Score a single user based on weighted criteria"""
    score = 0
    
    if filters.get("location"):
        location_input = filters["location"]
        location_query = str(location_input[0]).lower() if isinstance(location_input, list) and location_input else str(location_input).lower()
        if location_query:
            address = user.get("address", "").lower()
            user_licenses = [str(s).lower() for s in user.get("licenseState", [])]
            match = re.search(r"(?P<city>[\w\s.-]+?)\s*,\s*(?P<state>[a-z]{2})\b", address, re.IGNORECASE)
            if match:
                address_parts = {k: v.strip() for k, v in match.groupdict().items()}
                addr_city, addr_state_abbr = address_parts.get("city", ""), address_parts.get("state", "")
                addr_state_full = STATE_MAP.get(addr_state_abbr, "")
                if location_query in addr_city and addr_state_full in user_licenses:
                    score += 3
            if location_query in user_licenses:
                score += 2
    
    if filters.get("name") and user.get("name", "").lower() == filters["name"].lower():
        score += 10
    if filters.get("firstName") and user.get("firstName", "").lower() == filters["firstName"].lower():
        score += 5
    if filters.get("lastName") and user.get("lastName", "").lower() == filters["lastName"].lower():
        score += 5
    
    if filters.get("rating"):
        try:
            if float(filters["rating"]) <= float(user.get("rating", 0)) < float(filters["rating"]) + 1:
                score += 2
        except (ValueError, TypeError):
            pass

    if filters.get("hasCalendarConnected") and user.get("hasCalendarConnected"):
        score += 2
    
    if filters.get("meetingTypes"):
        filter_types = filters["meetingTypes"]
        if not isinstance(filter_types, list):
            filter_types = [filter_types]
        if any(m.lower() in [t.lower() for t in user.get("meetingTypes", [])] for m in filter_types):
            score += 1
    
    if filters.get("firm") and filters["firm"].lower() in user.get("firm", "").lower():
        score += 5

    if filters.get("languages"):
        filter_langs = filters["languages"]
        if not isinstance(filter_langs, list):
            filter_langs = [filter_langs]
        if any(language.lower() in [ul.lower() for ul in user.get("languages", [])] for language in filter_langs):
            score += 2

    if filters.get("review_keywords"):
        keywords = filters["review_keywords"]
        if not isinstance(keywords, list):
            keywords = [keywords]
        review_content = user.get("reviewContent", "").lower()
        for keyword in keywords:
            if keyword.lower() in review_content:
                score += 3
            
    return score

def explain_top_match(user_query: str, top_lawyer: dict, filters: dict):
    """Explain why a lawyer is a good match"""
    fallback_applied = filters.get("fallback_applied", False)
    prompt = ""
    if fallback_applied:
        original_specialty = filters.get("original_specialties", ["the requested field"])[0]
        fallback_specialty = filters.get("specialties", ["a related field"])[0]
        prompt = f"""A user wrote: "{user_query}"
No exact matches were found for "{original_specialty}". Explain why a lawyer who specializes in the related field of "{fallback_specialty}" is a good alternative.
You matched this lawyer:
Name: {top_lawyer.get('name', 'N/A')}, State: {top_lawyer.get('licenseState', 'N/A')}, Specialty: {top_lawyer.get('specialties', 'N/A')}, Rating: {top_lawyer.get('rating', 'N/A')}
Explain in 1–2 sentences why this lawyer is a good alternative match. Emphasize the connection between the original query and the lawyer's actual specialty. Be helpful and natural."""
    else:
        prompt = f"""A user wrote: "{user_query}"  
You matched this lawyer:
Name: {top_lawyer.get('name', 'N/A')}, State: {top_lawyer.get('licenseState', 'N/A')}, Specialty: {top_lawyer.get('specialties', 'N/A')}, Rating: {top_lawyer.get('rating', 'N/A')}
Explain in 1 sentence why this lawyer is the best match for the user query. Don't repeat exact attributes. Be helpful and natural. Dont mention "user", talk in 2nd person."""
    try:
        response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        return response.text.strip().strip("`")
    except Exception:
        return "This lawyer is a good match for your needs."

def handle_got_letter_flow(user_query: str, flow_context: dict, history: list):
    """Smart Gemini-powered handler for the got-letter flow"""
    step = flow_context.get('step', 1)
    data = flow_context.get('data', {})
    
    # Create a comprehensive prompt for Gemini to handle the flow intelligently
    prompt = f"""
You are handling a legal case intake flow for someone who received a letter from OpenLaw. 

Current step: {step}
Current data collected: {data}
User's latest message: "{user_query}"

Flow steps:
1. Reference number (from URL like https://olaw.io/REF123)
2. Ideal outcome for the case
3. Urgency level (high/medium/low)
4. Preferred language (English/Spanish)
5. Contact details (name, email, phone)
6. Mobile OTP verification (accept any 4-6 digit code)
7. Completion. 

Rules:
- If user wants to exit the flow, allow it gracefully with a professional response
- If the user goes off-topic, answer politely but briefly and then redirect them back to the flow.
- If user provides multiple pieces of information at once, extract all details that you can and only ask for remaining details
- Be conversational, professional, and slightly witty
- Validate inputs appropriately (email format, phone format, etc.)
- For OTP step, accept any 4-6 digit number as valid
- After OTP verification, end the flow with a completion message
- IMPORTANT: If user clearly states their desired outcome (like "dismissal", "win", "settlement", etc.), accept it and move to next step
- IMPORTANT: If user provides case number in any format, extract and accept it
- IMPORTANT: If user provides personal details (name, language preference), extract and store them

Current step details:
- Step 1: Ask for reference number from the letter URL (extract case numbers like AYU166, REF123, etc.)
- Step 2: Ask for ideal outcome they want to achieve (accept clear outcomes like "dismissal", "win", "settlement")
- Step 3: Ask for urgency level
- Step 4: Ask for preferred language
- Step 5: Ask for full name, email, and phone number
- Step 6: Ask for mobile OTP (accept any code)
- Step 7: Complete the flow

IMPORTANT: You must respond with ONLY a valid JSON object. No additional text before or after the JSON.

Respond with this exact JSON format:
{{
    "response": "your response to the user",
    "next_step": next_step_number,
    "extracted_data": {{"field": "value"}},
    "should_end_flow": true/false,
    "completion_message": "final message if ending flow"
}}

Be smart about extracting information and handling edge cases. Ensure the JSON is properly formatted.
"""
    
    try:
        response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        response_text = response.text.strip()
        
        # Try to extract JSON from the response
        try:
            # Look for JSON in the response
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}') + 1
            if start_idx != -1 and end_idx != 0:
                json_str = response_text[start_idx:end_idx]
                result = json.loads(json_str)
            else:
                raise ValueError("No JSON found in response")
                
        except (json.JSONDecodeError, ValueError) as json_error:
            print(f"JSON parsing error: {json_error}")
            print(f"Raw response: {response_text}")
            # Fallback to simple flow
            return handle_simple_got_letter_flow(user_query, flow_context)
        
        # Update data with extracted information
        if 'extracted_data' in result:
            data.update(result['extracted_data'])
        
        # Update flow context
        flow_context['step'] = result.get('next_step', step)
        flow_context['data'] = data
        
        return {
            'response': result['response'],
            'should_end_flow': result.get('should_end_flow', False),
            'completion_message': result.get('completion_message', ''),
            'flow_context': flow_context
        }
    except Exception as e:
        # Fallback to simple flow if Gemini fails
        print(f"Error in smart flow handler: {e}")
        return handle_simple_got_letter_flow(user_query, flow_context)

def handle_simple_got_letter_flow(user_query: str, flow_context: dict):
    """Fallback simple flow handler with smart features"""
    step = flow_context.get('step', 1)
    data = flow_context.get('data', {})
    
    # Check if user wants to exit the flow
    exit_phrases = ['exit', 'stop', 'cancel', 'quit', 'no thanks', 'not now', 'later']
    if any(phrase in user_query.lower() for phrase in exit_phrases):
        return {
            'response': "No problem! You can always come back to complete your case setup later. Is there anything else I can help you with regarding legal matters?",
            'should_end_flow': True,
            'flow_context': {'step': 0, 'data': data}
        }
    
    # Check if user is diverting from the flow
    diversion_phrases = ['weather', 'joke', 'how are you', 'what time', 'unrelated']
    if any(phrase in user_query.lower() for phrase in diversion_phrases):
        witty_responses = [
            "I appreciate the small talk, but let's focus on your legal case! 😊 ",
            "While I'd love to chat about that, your case needs attention! ",
            "Interesting topic, but your legal matter is more pressing! "
        ]
        import random
        witty_response = random.choice(witty_responses)
        
        if step == 1:
            return {
                'response': witty_response + "Could you please provide the reference number from your OpenLaw letter?",
                'should_end_flow': False,
                'flow_context': {'step': 1, 'data': data}
            }
        elif step == 2:
            return {
                'response': witty_response + "What's the ideal outcome you want to achieve for your case?",
                'should_end_flow': False,
                'flow_context': {'step': 2, 'data': data}
            }
        elif step == 3:
            return {
                'response': witty_response + "What is the level of urgency for your case (e.g., high, medium, low)?",
                'should_end_flow': False,
                'flow_context': {'step': 3, 'data': data}
            }
        elif step == 4:
            return {
                'response': witty_response + "What is your preferred language, English or Spanish?",
                'should_end_flow': False,
                'flow_context': {'step': 4, 'data': data}
            }
        elif step == 5:
            return {
                'response': witty_response + "Could you please provide your full name, email, and phone number?",
                'should_end_flow': False,
                'flow_context': {'step': 5, 'data': data}
            }
        elif step == 6:
            return {
                'response': witty_response + "Please enter the 4-digit OTP sent to your mobile number:",
                'should_end_flow': False,
                'flow_context': {'step': 6, 'data': data}
            }
    
    # Handle multiple inputs at once
    if step == 5 and ('@' in user_query and any(char.isdigit() for char in user_query)):
        # User likely provided contact details
        data['contact_details'] = user_query
        return {
            'response': "Please enter the 4-digit OTP sent to your mobile number:",
            'should_end_flow': False,
            'flow_context': {'step': 6, 'data': data}
        }
    
    # Normal flow progression
    if step == 1:
        # Extract case number from various formats
        import re
        case_pattern = r'\b[A-Z]{2,4}\d{3,6}\b|\bREF\d{3,6}\b'
        case_match = re.search(case_pattern, user_query.upper())
        if case_match:
            data['reference_number'] = case_match.group()
        else:
            data['reference_number'] = user_query
        
        # Check if user also provided outcome in the same message
        outcome_keywords = ['dismiss', 'dismissal', 'win', 'settle', 'settlement', 'resolve', 'drop', 'withdraw']
        if any(keyword in user_query.lower() for keyword in outcome_keywords):
            for keyword in outcome_keywords:
                if keyword in user_query.lower():
                    data['outcome'] = keyword
                    return {
                        'response': "What is the level of urgency for your case (e.g., high, medium, low)?",
                        'should_end_flow': False,
                        'flow_context': {'step': 3, 'data': data}
                    }
        
        return {
            'response': "What's the ideal outcome you want to achieve for your case?",
            'should_end_flow': False,
            'flow_context': {'step': 2, 'data': data}
        }
    elif step == 2:
        # Extract outcome from user response
        outcome_keywords = ['dismiss', 'dismissal', 'win', 'settle', 'settlement', 'resolve', 'drop', 'withdraw']
        extracted_outcome = None
        for keyword in outcome_keywords:
            if keyword in user_query.lower():
                extracted_outcome = keyword
                break
        
        if extracted_outcome:
            data['outcome'] = extracted_outcome
        else:
            data['outcome'] = user_query
            
        return {
            'response': "What is the level of urgency for your case (e.g., high, medium, low)?",
            'should_end_flow': False,
            'flow_context': {'step': 3, 'data': data}
        }
    elif step == 3:
        data['urgency'] = user_query
        return {
            'response': "What is your preferred language, English or Spanish?",
            'should_end_flow': False,
            'flow_context': {'step': 4, 'data': data}
        }
    elif step == 4:
        data['language'] = user_query
        return {
            'response': "Could you please provide your full name, email, and phone number?",
            'should_end_flow': False,
            'flow_context': {'step': 5, 'data': data}
        }
    elif step == 5:
        data['contact_details'] = user_query
        return {
            'response': "Please enter the 4-digit OTP sent to your mobile number:",
            'should_end_flow': False,
            'flow_context': {'step': 6, 'data': data}
        }
    elif step == 6:
        data['otp'] = user_query
        return {
            'response': "Perfect! Your case has been successfully registered. Click here to check your newly created profile.",
            'should_end_flow': True,
            'completion_message': "🎉 Welcome to OpenLaw! Your case profile has been created successfully. Our team will review your information and get back to you within 24 hours. You can track your case status in your new profile dashboard.",
            'flow_context': {'step': 7, 'data': data}
        }

def handle_document_drafting_flow(user_query: str, flow_context: dict, history: list):
    """Intelligent LLM-powered document drafting flow"""
    step = flow_context.get('step', 1)
    data = flow_context.get('data', {})
    
    # Step 1: Initial document type determination
    if step == 1:
        requirements = determine_document_requirements(user_query, history)
        data.update(requirements.get('extracted_data', {}))
        data['document_type'] = requirements.get('document_type', 'legal_document')
        
        flow_context['step'] = 2
        flow_context['data'] = data
        
        return {
            'response': requirements.get('next_question', 'What type of legal document do you need to create?'),
            'should_end_flow': False,
            'flow_context': flow_context
        }
    
    # Step 2: Gather additional information
    elif step == 2:
        # Use LLM to extract information and determine next question
        prompt = f"""
You are helping gather information for a legal document. 

Document Type: {data.get('document_type', 'legal document')}
Current Data: {json.dumps(data, indent=2)}
User's Response: "{user_query}"

Analyze the user's response and:
1. Extract any relevant information provided
2. Determine what additional information is still needed
3. Ask the next most important question
4. Wrap up the flow in at max 12-15 questions. these questions should give you enough information to generate the document. 
5. Follow standard legal document drafting practices and formats.
6. If the user provides multiple pieces of information, extract all you can.
7. If the user provides a lot of information, ask for clarification on any specific details.
8. If user wants a blueprint to get started, make it using whatever information you have.
9. If user wants to exit the flow, allow it gracefully with a professional response.
10. Before giving the user the document, always give a caution saying that the document is not a legal advice and is for informational purposes only, they should consult a lawyer for legal advice.

Respond with JSON:
{{
    "extracted_data": {{"field": "value"}},
    "next_question": "your next question",
    "has_sufficient_info": true/false,
    "missing_info": ["list of still needed information"]
}}

Be conversational and professional. If the user provides multiple pieces of information, extract all you can.
"""
        
        try:
            response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
            response_text = response.text.strip()
            
            # Extract JSON
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}') + 1
            if start_idx != -1 and end_idx != 0:
                result = json.loads(response_text[start_idx:end_idx])
            else:
                raise ValueError("No JSON found")
            
            # Update data
            data.update(result.get('extracted_data', {}))
            
            # Check if we have sufficient information
            if result.get('has_sufficient_info', False):
                # Generate the document
                document_content = generate_legal_document(data, history)
                flow_context['step'] = 3
                
                return {
                    'response': "Perfect! I have enough information to create your document. Let me generate it for you.",
                    'should_end_flow': True,
                    'document_content': document_content,
                    'document_type': data.get('document_type', 'legal_document'),
                    'flow_context': flow_context
                }
            else:
                # Continue gathering information
                flow_context['step'] = 2
                flow_context['data'] = data
                
                return {
                    'response': result.get('next_question', 'Could you provide more details?'),
                    'should_end_flow': False,
                    'flow_context': flow_context
                }
                
        except Exception as e:
            print(f"Error in document flow: {e}")
            # Fallback to simple flow
            return handle_simple_drafting_flow(user_query, flow_context)
    
    # Step 3: Document generation (shouldn't normally reach here)
    else:
        document_content = generate_legal_document(data, history)
        return {
            'response': "Here's your generated document!",
            'should_end_flow': True,
            'document_content': document_content,
            'document_type': data.get('document_type', 'legal_document'),
            'flow_context': flow_context
        }

def handle_simple_drafting_flow(user_query: str, flow_context: dict):
    """Fallback simple document drafting flow handler"""
    step = flow_context.get('step', 1)
    data = flow_context.get('data', {})
    
    # Check if user wants to proceed with standard terms
    if any(phrase in user_query.lower() for phrase in ['standard', 'no', 'not really', 'use standard', 'default']):
        # User wants to use standard terms, generate document
        document_content = generate_legal_document(data)
        return {
            'response': "Perfect! I've drafted your document using standard terms and conditions. You can download it below:",
            'should_end_flow': True,
            'document_content': document_content,
            'document_type': data.get('document_type', 'legal_document'),
            'flow_context': {'step': 4, 'data': data}
        }
    
    if step == 1:
        data['document_type'] = user_query
        return {
            'response': "Great! Now I need some key details. What is the purpose of this document? Who are the parties involved?",
            'should_end_flow': False,
            'flow_context': {'step': 2, 'data': data}
        }
    elif step == 2:
        data['purpose'] = user_query
        return {
            'response': "Perfect! Now let me ask about the specific terms or conditions you want to include in this document. You can say 'standard' if you want to use default terms.",
            'should_end_flow': False,
            'flow_context': {'step': 3, 'data': data}
        }
    elif step == 3:
        data['terms'] = user_query
        # Generate the document
        document_content = generate_legal_document(data)
        return {
            'response': "I've drafted your document! You can download it below:",
            'should_end_flow': True,
            'document_content': document_content,
            'document_type': data.get('document_type', 'legal_document'),
            'flow_context': {'step': 4, 'data': data}
        }

def generate_legal_document(data: dict, conversation_history: list = None):
    """Generate a legal document using LLM based on collected data and conversation context"""
    
    # Build context from conversation history
    conversation_context = ""
    if conversation_history:
        recent_messages = conversation_history[-25:]  # Last 25 messages for context
        conversation_context = "\n".join([
            f"{msg['role']}: {msg['parts'][0]['text']}" 
            for msg in recent_messages
        ])
    
    # Create a comprehensive prompt for document generation
    prompt = f"""
You are a legal document drafting expert. Generate a professional, legally sound document based on the provided information.

Document Data: {json.dumps(data, indent=2)}

Conversation Context:
{conversation_context}

Instructions:
1. Analyze the document type and requirements from the data and conversation
2. Determine what type of legal document is needed (contract, agreement, letter, notice, etc.)
3. Identify all parties involved and their roles
4. Extract all relevant details, terms, conditions, and requirements
5. Generate a complete, professional legal document that includes:
   - Proper legal formatting and structure
   - Clear, unambiguous language
   - All necessary sections and clauses for the document type
   - Appropriate legal disclaimers and boilerplate
   - Standard legal language and terminology
   - All parties' information and signatures
   - Date, jurisdiction, and legal references

Document Requirements:
- Make it ready for immediate use as a legal document
- Include all standard clauses and sections for the document type
- Use proper legal formatting and structure
- Ensure all terms are clear and enforceable
- Include appropriate legal disclaimers
- Make it comprehensive and professional

Generate the complete legal document:
"""
    
    try:
        response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error generating document: {e}"

def determine_document_requirements(user_query: str, conversation_history: list = None):
    """Use LLM to intelligently determine what information is needed for document creation"""
    
    conversation_context = ""
    if conversation_history:
        recent_messages = conversation_history[-20:]  # Last 20 messages
        conversation_context = "\n".join([
            f"{msg['role']}: {msg['parts'][0]['text']}" 
            for msg in recent_messages
        ])
    
    prompt = f"""
You are a legal assistant helping a user create a legal document. Analyze the user's request and determine what information is needed.

User Query: "{user_query}"

Conversation Context:
{conversation_context}

Your task:
1. Determine what type of legal document the user needs
2. Identify what specific information is required to create this document
3. Ask intelligent, targeted questions to gather the necessary details
4. Be conversational but professional
5. If the user provides multiple pieces of information, extract what you can and ask for the remaining details

Respond with a JSON object in this exact format:
{{
    "document_type": "type of document (e.g., rental agreement, employment contract, cease and desist letter)",
    "required_info": ["list of specific information needed"],
    "next_question": "your next question to the user",
    "extracted_data": {{"field": "value"}},
    "confidence": "high/medium/low"
}}

Be specific about what information is needed for the document type.
"""
    
    try:
        response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        response_text = response.text.strip()
        
        # Extract JSON from response
        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1
        if start_idx != -1 and end_idx != 0:
            json_str = response_text[start_idx:end_idx]
            return json.loads(json_str)
        else:
            raise ValueError("No JSON found in response")
            
    except Exception as e:
        return {
            "document_type": "legal document",
            "required_info": ["document type", "parties involved", "terms and conditions"],
            "next_question": "What type of legal document do you need to create?",
            "extracted_data": {},
            "confidence": "low"
        }


def create_docx(content, filename):
    """Create a .docx file from content"""
    doc = Document()
    
    # Split content into paragraphs and add to document
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Handle markdown formatting
            if para.startswith('**') and para.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                p.add_run(para[2:-2]).bold = True
            elif para.startswith('#'):
                # Heading
                level = len(para) - len(para.lstrip('#'))
                text = para.lstrip('#').strip()
                if level == 1:
                    doc.add_heading(text, level=1)
                elif level == 2:
                    doc.add_heading(text, level=2)
                else:
                    doc.add_heading(text, level=3)
            else:
                # Regular paragraph
                doc.add_paragraph(para.strip())
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Firebase is used for chat session storage
from flask import request, jsonify
import tempfile, os, uuid, json
from google import genai
from google.genai import types

# Assume `app`, `client`, `db`, `firestore`, and all handler functions are already defined

@app.route('/search', methods=['POST'])
def search_users():
    """Main search endpoint"""
    try:
        query = request.form.get('query')
        chat_id = request.form.get('chat_id')
        file = request.files.get('file')

        # Initialize or get chat session from Firebase
        history = []
        flow_context = {}
        if chat_id and db:
            try:
                chat_ref = db.collection('chats').document(chat_id)
                chat_doc = chat_ref.get()
                if chat_doc.exists:
                    chat_data = chat_doc.to_dict()
                    history = chat_data.get('history', [])
                    flow_context = chat_data.get('flow_context', {})
            except Exception as e:
                print(f"Error accessing Firebase: {e}")

        # Handle active flow sessions
        if flow_context.get('name') == 'got-letter':
            flow_result = handle_got_letter_flow(query, flow_context, history)
            response_text = flow_result['response']
            should_end_flow = flow_result['should_end_flow']
            completion_message = flow_result.get('completion_message', '')
            response_data = {"answer": response_text}

            history.append({"role": "user", "parts": [{"text": query}]})
            history.append({"role": "model", "parts": [{"text": response_text}]})
            if should_end_flow and completion_message:
                history.append({"role": "model", "parts": [{"text": completion_message}]})
                response_data["completion_message"] = completion_message

            if db:
                try:
                    db_update = {'history': history}
                    db_update['flow_context'] = firestore.DELETE_FIELD if should_end_flow else flow_result['flow_context']
                    db.collection('chats').document(chat_id).set(db_update, merge=True)
                except Exception as e:
                    print(f"Error saving to Firebase: {e}")

            response_data['chat_id'] = chat_id
            return jsonify(response_data)

        elif flow_context.get('name') == 'draft_document':
            flow_result = handle_document_drafting_flow(query, flow_context, history)
            response_text = flow_result['response']
            should_end_flow = flow_result['should_end_flow']
            document_content = flow_result.get('document_content', '')
            document_type = flow_result.get('document_type', 'legal_document')

            response_data = {"answer": response_text}
            history.append({"role": "user", "parts": [{"text": query}]})
            history.append({"role": "model", "parts": [{"text": response_text}]})

            if should_end_flow and document_content:
                response_data.update({
                    "document_content": document_content,
                    "document_type": document_type,
                    "download_filename": f"{document_type.replace(' ', '_')}.docx",
                    "answer": "Perfect! I've drafted your document. You can download it below:"
                })

            if db:
                try:
                    db_update = {'history': history}
                    db_update['flow_context'] = firestore.DELETE_FIELD if should_end_flow else flow_result['flow_context']
                    db.collection('chats').document(chat_id).set(db_update, merge=True)
                except Exception as e:
                    print(f"Error saving to Firebase: {e}")

            response_data['chat_id'] = chat_id
            return jsonify(response_data)

        # Handle file-only uploads (document analysis)
        if query is None and file:
            query = "Please analyze this legal document"

        if file:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name

            try:
                prompt = (
                    "You are a legal assistant and are here to give legal advice. "
                    "Do not say you're an AI model. Read the document and determine:\n"
                    "- Is it a legal document (contract, notice, policy, court order, etc)? If yes, explain the document in 2–5 sentences."
                    "If it is clearly not a legal document, respond only with:\n"
                    "\"That does not look like a legal document, please upload a legal document.\""
                )

                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=[
                        types.Part.from_bytes(
                            data=open(tmp_path, "rb").read(),
                            mime_type='application/pdf'
                        ),
                        prompt
                    ]
                )

                analysis_result = response.text.strip().strip("`")
                response_data = {"answer": analysis_result}

                if not chat_id:
                    chat_id = str(uuid.uuid4())
                response_data['chat_id'] = chat_id

                history.append({"role": "user", "parts": [{"text": f"Uploaded document: {file.filename}"}]})
                history.append({"role": "model", "parts": [{"text": analysis_result}]})

                if db:
                    try:
                        db.collection('chats').document(chat_id).set({'history': history}, merge=True)
                    except Exception as e:
                        print(f"Error saving to Firebase: {e}")

                return jsonify(response_data)

            except Exception as e:
                return jsonify({"error": f"Failed to analyze file: {str(e)}"}), 500
            finally:
                os.unlink(tmp_path)

        # Handle pure text-based queries (intent classification, matching, flows, etc.)
        if not query:
            return jsonify({"error": "No query or file provided"}), 400

        user_message = {"role": "user", "parts": [{"text": query}]}
        history.append(user_message)

        parsed = parse_query_with_gemini(history)
        intent = parsed.get("intent")

        # Handle flows from scratch
        if intent == "got-letter":
            if any(word in query.lower() for word in ['ref', 'reference', 'olaw.io', 'https://']):
                new_flow_context = {'name': 'got-letter', 'step': 1, 'data': {}}
                flow_result = handle_got_letter_flow(query, new_flow_context, history)
                response_text = flow_result['response']
                should_end_flow = flow_result['should_end_flow']
                completion_message = flow_result.get('completion_message', '')
            else:
                response_text = "I can help you with your OpenLaw case! To get started, could you please provide the reference number from your letter? You can find it in the URL, like this: https://olaw.io/YOUR_REFERENCE_NUMBER."
                should_end_flow = False
                completion_message = ""
                new_flow_context = {'name': 'got-letter', 'step': 1, 'data': {}}

            response_data = {"answer": response_text}
            if not chat_id:
                chat_id = str(uuid.uuid4())
            history.append({"role": "model", "parts": [{"text": response_text}]})

            if should_end_flow and completion_message:
                history.append({"role": "model", "parts": [{"text": completion_message}]})
                response_data["completion_message"] = completion_message

            if db:
                try:
                    db_update = {'history': history}
                    db_update['flow_context'] = firestore.DELETE_FIELD if should_end_flow else new_flow_context
                    db.collection('chats').document(chat_id).set(db_update, merge=True)
                except Exception as e:
                    print(f"Error saving to Firebase: {e}")

            response_data['chat_id'] = chat_id
            return jsonify(response_data)

        elif intent == "draft_document":
            response_text = "I can help you draft a legal document! To create the most appropriate document for your needs, I'll need some details. What type of legal document are you looking to create? (e.g., contract, letter, notice, agreement, etc.)"
            response_data = {"answer": response_text}
            if not chat_id:
                chat_id = str(uuid.uuid4())

            new_flow_context = {'name': 'draft_document', 'step': 1, 'data': {}}
            history.append({"role": "model", "parts": [{"text": response_text}]})

            if db:
                try:
                    db.collection('chats').document(chat_id).set({'history': history, 'flow_context': new_flow_context}, merge=True)
                except Exception as e:
                    print(f"Error saving to Firebase: {e}")

            response_data['chat_id'] = chat_id
            return jsonify(response_data)

        if intent == "near_me":
            return jsonify({"intent": "near_me"})

        intent_responses = {
            "general_question": {
                "answer": parsed.get("response"),
                "message": "That sounds like a legal question. I can help you find a lawyer — Can you describe your case or what kind of lawyer you need and where."
            },
            "out-of-scope": {
                "answer": parsed.get("witty_response", "I am having trouble understanding that, can you please rephrase? But I can help you with legal matters!")
            },
            "greeting": {
                "answer": "Hi there! How can I help you?"
            },
            "out-of-area": {
                "answer": "We're sorry, but we don't currently operate in your state. We're expanding and hope to be available in your area soon!",
                "show_form": True
            },
            "document_help": {
                "answer": "I understand you need help understanding a legal document. While I can provide general information and help you understand the document's content, it's important to note that for specific legal advice, you should consult with a qualified attorney who can review your unique situation.",
                "disclaimer": "⚠️ **Important Disclaimer**: This AI assistance is for informational purposes only and should not be considered legal advice. For specific legal guidance, please consult with a qualified attorney.",
                "encourage_upload": "To help you better, please upload your legal document using the file upload feature. I can then analyze the document and provide you with a general explanation of its contents, key terms, and what it might mean in plain language."
            },
            "goodbye": {
                "answer": "generate_goodbye_message"
            }
        }

        if intent in intent_responses:
            if intent == "goodbye":
                # Generate dynamic goodbye message
                goodbye_message = generate_goodbye_message(history)
                response_data = {"query": query, "answer": goodbye_message}
            else:
                response_data = {"query": query, **intent_responses[intent]}
        elif "error" in parsed:
            response_data = parsed
        else:
            filters = {k: v for k, v in parsed.items() if v}
            top_matches = find_best_matches(filters)

            if not top_matches and "specialties" in filters:
                original_specialties = filters["specialties"]
                if not isinstance(original_specialties, list):
                    original_specialties = [original_specialties]
                fallback_specialties = [spec for s in original_specialties for spec in SPECIALTY_FALLBACK_MAP.get(s.lower(), [])]
                if fallback_specialties:
                    fallback_filters = filters.copy()
                    fallback_filters["specialties"] = list(set(fallback_specialties))
                    top_matches = find_best_matches(fallback_filters)
                    if top_matches:
                        filters["fallback_applied"] = True
                        filters["original_specialties"] = original_specialties

            if top_matches:
                for lawyer in top_matches:
                    lawyer['explanation'] = explain_top_match(query, lawyer, filters)

            response_data = {"query": query, "filters_applied": filters, "matches": top_matches}

        ai_response_text = json.dumps(response_data)
        history.append({"role": "model", "parts": [{"text": ai_response_text}]})

        if not chat_id:
            chat_id = str(uuid.uuid4())
        if db:
            try:
                db.collection('chats').document(chat_id).set({'history': history}, merge=True)
            except Exception as e:
                print(f"Error saving to Firebase: {e}")

        response_data['chat_id'] = chat_id
        return jsonify(response_data)

    except Exception as e:
        return jsonify({"error": f"Unexpected error: {str(e)}"}), 500


@app.route('/generate-document', methods=['POST'])
def generate_document():
    """Generate and return a .docx document"""
    try:
        document_content = request.form.get('document_content')
        document_type = request.form.get('document_type', 'Document')
        filename = request.form.get('filename', 'document.docx')
        
        if not document_content:
            return jsonify({"error": "No document content provided"}), 400
        
        # Create the document
        buffer = create_docx(document_content, filename)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload for document analysis"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        # Save the uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        # Use Gemini to analyze the document
        prompt = f"""You are a legal assistant and are here to give legal advice. Dont ever say you are an AI Model or cant help with legal advice. Read the document and determine:
                - Is it a legal document (contract, notice, policy, court order, etc)?
                - If yes, explain the document in 2–5 sentences.

                If it is clearly not a legal document, respond only with:
                "That does not look like a legal document, please upload a legal document."""
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[
                types.Part.from_bytes(
                    data=open(tmp_path, "rb").read(),
                    mime_type='application/pdf'
                ),
                prompt
            ]
        )

        # Clean up temporary file
        os.unlink(tmp_path)

        return jsonify({"answer": response.text.strip().strip("`")})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/generate-title', methods=['POST'])
def generate_title():
    """Generate a title for a chat conversation"""
    try:
        data = request.get_json()
        conversation_text = data.get("conversation", "")
        
        prompt = f"""
Based on this conversation, generate a short, descriptive title (maximum 50 characters) that captures the main topic or legal issue being discussed. The title should be clear and helpful for identifying the chat later.

Conversation:
{conversation_text}

Generate only the title, nothing else:
"""
        
        response = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        title = response.text.strip().strip('"').strip("'")
        
        # Ensure title is not too long
        if len(title) > 50:
            title = title[:47] + "..."
        
        return jsonify({"title": title})
        
    except Exception as e:
        return jsonify({"error": f"Error generating title: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8000) 